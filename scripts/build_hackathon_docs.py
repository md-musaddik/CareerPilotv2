from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"

PRIMARY = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT = RGBColor(0xF3, 0xF6, 0xFA)
BORDER = RGBColor(0xD8, 0xE1, 0xEB)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D8E1EB") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, PRIMARY, 16, 8),
        ("Heading 2", 13, PRIMARY, 12, 6),
        ("Heading 3", 12, PRIMARY, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(short_title)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("CareerPilot Hackathon Submission")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = PRIMARY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2.font.size = Pt(11)
    r2.font.color.rgb = MUTED


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.85)
    table.columns[1].width = Inches(4.65)

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.85)
        cells[1].width = Inches(4.65)
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], "F3F6FA")
        p_label = cells[0].paragraphs[0]
        p_value = cells[1].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(0)
        p_value.paragraph_format.space_after = Pt(0)
        rl = p_label.add_run(label)
        rl.font.bold = True
        rl.font.color.rgb = TEXT
        rv = p_value.add_run(value)
        rv.font.color.rgb = TEXT


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].width = Inches(widths[idx])
        set_cell_shading(header_cells[idx], "E8EEF5")
        set_cell_border(header_cells[idx], "C5D3E0")
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = header_cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.font.bold = True
        r.font.color.rgb = TEXT
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)


def build_stack_report(path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, "Stack Report & Justification")
    add_title_block(
        doc,
        "Stack Report & Justification",
        "CareerPilot | Hackathon submission support document",
    )

    intro = doc.add_paragraph()
    intro.add_run(
        "This report explains the chosen stack for CareerPilot, why each major platform decision fits the problem statement, "
        "and what tradeoffs the team intentionally accepted to ship a working end-to-end product during the hackathon."
    )

    doc.add_heading("Executive Summary", level=1)
    add_key_value_table(
        doc,
        [
            ("Product", "CareerPilot"),
            ("Core goal", "Turn a user's CV into the source of truth for search, fit scoring, AI guidance, and productivity tracking."),
            ("Frontend", "React + Vite + TypeScript"),
            ("Backend", "Node.js + Express + TypeScript"),
            ("Database", "MongoDB Atlas + Atlas Vector Search"),
            ("Auth", "Firebase Auth"),
            ("Storage", "Cloudinary for original CV files"),
            ("AI", "OpenAI for embeddings and assistant generation"),
            ("External tools", "Adzuna + Jooble for live job search"),
        ],
    )

    doc.add_heading("Why This Stack Fits The Problem Statement", level=1)
    add_bullets(
        doc,
        [
            "The CV can be parsed, chunked, embedded, and reused across agents without building separate profile stores for every feature.",
            "The backend cleanly owns secret-bearing integrations such as OpenAI, Adzuna, Firebase Admin, and Cloudinary.",
            "MongoDB Atlas supports both product data and vector-search workflows, which keeps the architecture small enough for a hackathon while still feeling production-shaped.",
            "The chosen frontend stack supports a dashboard-style command center with fast iteration, strong type safety, and route-level organization.",
            "The result is a single platform that covers all four pillars instead of stitching together disconnected demos.",
        ],
    )

    doc.add_heading("Stack Decisions By Layer", level=1)
    add_matrix_table(
        doc,
        ["Layer", "Choice", "Justification"],
        [
            [
                "Frontend",
                "React, Vite, TypeScript, React Router, Tailwind, shadcn/ui, React Query",
                "This combination supports a fast dashboard-like product, typed contracts, reusable UI primitives, and predictable async state for jobs, resume workflows, and tracker views.",
            ],
            [
                "Backend",
                "Node.js, Express, TypeScript",
                "The API needed to orchestrate file uploads, parsing, embeddings, vector retrieval, fit scoring, and streaming chat. Express keeps the HTTP layer simple while TypeScript keeps the service boundaries explicit.",
            ],
            [
                "Database",
                "MongoDB Atlas",
                "CareerPilot stores user-scoped documents such as resumes, profile state, applications, goals, tasks, calendar events, and chat history. MongoDB fits this document-heavy model well.",
            ],
            [
                "Vector search",
                "MongoDB Atlas Vector Search",
                "Using Atlas Vector Search avoided a second infrastructure surface for embeddings and kept the RAG pipeline close to the existing product data.",
            ],
            [
                "Authentication",
                "Firebase Auth",
                "Firebase provides a fast path to secure email and Google sign-in, while backend token verification keeps user-owned data scoped safely by UID.",
            ],
            [
                "File storage",
                "Cloudinary",
                "Cloudinary replaced Firebase Storage because the project needed a practical resume-file host without relying on a premium Firebase storage path.",
            ],
            [
                "AI",
                "OpenAI",
                "OpenAI is used in a bounded way: embeddings for RAG, assistant generation for grounded responses, and optional explanation text after deterministic scoring already exists.",
            ],
            [
                "Jobs sources",
                "Adzuna API + Jooble API",
                "Adzuna covers the configured US market while Jooble broadens international search coverage. The backend merges, filters, dedupes, and ranks results after parsing one natural-language query.",
            ],
            [
                "Deployment",
                "Vercel + Railway",
                "This split keeps the frontend and backend deploy surfaces simple, works well with Git-based deployment, and matches the team's local development flow.",
            ],
        ],
        [1.0, 1.8, 3.7],
    )

    doc.add_heading("How The Stack Satisfies Core Technical Requirements", level=1)
    add_numbered(
        doc,
        [
            "RAG grounded in the user's actual CV is supported by resume parsing, chunking, OpenAI embeddings, and MongoDB Atlas Vector Search.",
            "At least one agent uses external tool calls through the live Adzuna and Jooble job-search integrations.",
            "Conversational memory within a session is supported through stored chat session state and restored frontend session context.",
            "Fit score is computed programmatically through deterministic criteria instead of being invented by the LLM.",
            "The tracker module includes live applications, goals, tasks, and calendar views connected to backend workspace data.",
        ],
    )

    doc.add_heading("Key Tradeoffs And Justifications", level=1)
    add_matrix_table(
        doc,
        ["Tradeoff", "Why it was chosen", "Impact"],
        [
            [
                "MongoDB for both app data and vectors",
                "This reduced infrastructure sprawl and kept the hackathon implementation focused.",
                "Simpler architecture, slightly less specialized than a dedicated vector database.",
            ],
            [
                "Backend-only AI orchestration",
                "Needed for security, grounding, and future portability of prompts and memory assembly.",
                "Slightly more backend complexity, much safer secret handling.",
            ],
            [
                "Cloudinary instead of Firebase Storage",
                "Practical response to storage-plan limits while preserving resume upload behavior.",
                "Small provider swap, no meaningful product regression.",
            ],
            [
                "Deterministic score plus AI explanation",
                "This keeps scoring inspectable and judge-friendly while still allowing natural-language explanation.",
                "Higher trust and explainability than a pure LLM-scored system.",
            ],
        ],
        [1.5, 2.3, 2.7],
    )

    doc.add_heading("Why The Team Would Keep This Stack After The Hackathon", level=1)
    add_bullets(
        doc,
        [
            "The separation between frontend, backend, and shared contracts is strong enough to keep extending safely.",
            "The data model already supports real tracker workflows instead of a one-time AI demo.",
            "The stack scales by adding backend instances and refining indexing and caching rather than rewriting the system.",
            "The architecture stays understandable for judges, teammates, and future maintainers.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        "The chosen stack is justified not because it is trendy, but because it supports grounded AI, deterministic scoring, live search, and real user workflow execution in one coherent product. "
        "It is small enough to ship during a hackathon and structured enough to evolve after the event."
    )

    doc.save(path)


def build_dependencies_doc(path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, "Dependencies & Documentation")
    add_title_block(
        doc,
        "Dependencies & Documentation",
        "CareerPilot | Runtime, libraries, and supporting project documents",
    )

    intro = doc.add_paragraph()
    intro.add_run(
        "This document inventories the major runtime dependencies used in CareerPilot and maps them to the project features they support. "
        "It also points reviewers to the internal documentation that explains how the product works."
    )

    doc.add_heading("Runtime Platforms", level=1)
    add_matrix_table(
        doc,
        ["Category", "Dependency", "Use in CareerPilot"],
        [
            ["Language/runtime", "Node.js", "Runs the frontend toolchain and the Express backend."],
            ["Language/runtime", "TypeScript", "Provides typed contracts across frontend, backend, and shared modules."],
            ["Frontend framework", "React", "Builds dashboard, jobs, AI assistant, and workspace interfaces."],
            ["Build tool", "Vite", "Runs local frontend development and production bundling."],
            ["Backend framework", "Express", "Hosts the API, auth middleware, and service orchestration."],
            ["Database", "MongoDB Atlas", "Stores user-scoped product data."],
            ["Vector search", "Atlas Vector Search", "Supports retrieval over embedded CV chunks."],
            ["Authentication", "Firebase Auth / Firebase Admin", "Handles sign-in on the client and token verification on the server."],
            ["File storage", "Cloudinary", "Stores original uploaded CV files."],
            ["AI provider", "OpenAI", "Provides embeddings and grounded assistant generation."],
            ["External jobs sources", "Adzuna API + Jooble API", "Supply live job-search results that the backend parses, merges, filters, and ranks."],
        ],
        [1.2, 1.8, 3.5],
    )

    doc.add_heading("Frontend Package Dependencies", level=1)
    add_matrix_table(
        doc,
        ["Package", "Purpose"],
        [
            ["react / react-dom", "Render the CareerPilot web application."],
            ["react-router-dom", "Provide route-level navigation between dashboard areas."],
            ["@tanstack/react-query", "Handle server-state fetching, caching, and mutation state."],
            ["firebase", "Support Firebase Auth flows in the client."],
            ["lucide-react", "Provide consistent UI icons."],
            ["tailwindcss", "Drive utility-first styling."],
            ["@radix-ui/*", "Support accessible primitive components used by shadcn/ui patterns."],
            ["class-variance-authority, clsx, tailwind-merge", "Manage conditional class composition."],
        ],
        [2.4, 4.1],
    )

    doc.add_heading("Backend Package Dependencies", level=1)
    add_matrix_table(
        doc,
        ["Package", "Purpose"],
        [
            ["express", "HTTP API layer."],
            ["mongoose", "MongoDB models and document access."],
            ["firebase-admin", "Verify Firebase ID tokens on the backend."],
            ["cloudinary", "Upload and reference original resume files."],
            ["multer", "Handle multipart file uploads."],
            ["pdf-parse", "Extract text from uploaded PDF resumes."],
            ["mammoth", "Extract text from uploaded DOCX resumes."],
            ["openai", "Create embeddings and assistant responses."],
            ["zod", "Validate request input and configuration boundaries."],
            ["cors", "Control browser access to backend routes."],
            ["dotenv", "Load runtime configuration from environment variables."],
        ],
        [2.2, 4.3],
    )

    doc.add_heading("Search And RAG-Specific Dependencies", level=1)
    add_matrix_table(
        doc,
        ["Capability", "Dependency", "Role"],
        [
            ["Resume extraction", "pdf-parse + mammoth", "Convert uploaded PDF and DOCX resumes into raw text."],
            ["Chunk storage", "MongoDB Atlas", "Persist parsed resumes, resume chunks, and associated metadata."],
            ["Vector retrieval", "Atlas Vector Search", "Retrieve the most relevant chunks for assistant grounding."],
            ["Embeddings", "OpenAI", "Embed section-based CV chunks with `text-embedding-3-small`."],
            ["US job source", "Adzuna API", "Live search provider for the configured US market."],
            ["International job source", "Jooble API", "Broader market coverage and fallback provider for international search."],
        ],
        [1.7, 1.8, 3.0],
    )

    doc.add_heading("Project Documentation Map", level=1)
    add_matrix_table(
        doc,
        ["Document", "What it covers"],
        [
            ["README.md", "Setup, required environment variables, local run steps, and repository overview."],
            ["docs/architecture/system-diagram.md", "End-to-end data flow from CV upload to grounded agent response."],
            ["docs/technical-documentation.md", "Major technical decisions and implementation notes."],
            ["docs/system-design-document.md", "Scaling, cost, bottlenecks, and architecture analysis."],
            ["docs/judge-demo-script.md", "Required 5-minute demo flow for the hackathon video."],
            ["docs/environment-variable-guide.md", "Environment variable definitions and usage boundaries."],
            ["docs/local-setup-guide.md", "Local development setup and validation steps."],
            ["docs/deployment-guide.md", "Deployment notes for Vercel and Railway."],
            ["docs/github-vercel-railway-workflow.md", "GitHub, branch, preview, and deployment workflow guidance."],
            ["AGENTS.md", "Permanent architectural source of truth for the project."],
        ],
        [2.5, 3.9],
    )

    doc.add_heading("Why These Dependencies Are Appropriate", level=1)
    add_bullets(
        doc,
        [
            "They are aligned to the hackathon problem rather than added for novelty.",
            "They support the required pillars without forcing unnecessary infrastructure.",
            "They keep AI secrets and third-party credentials on the backend only.",
            "They provide a clear path from prototype quality to a more production-shaped architecture.",
        ],
    )

    doc.add_heading("Operational Notes For Reviewers", level=1)
    add_bullets(
        doc,
        [
            "Frontend and backend are separate deploy units, which makes debugging and redeployment simpler.",
            "The shared package reduces drift between request and response contracts.",
            "The system relies on live third-party services, so the environment variables must be configured correctly before local or hosted testing.",
            "Resume upload, job search, fit scoring, and AI assistant flows should all be verified after deployment because they depend on external services.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The dependency set is intentionally practical: every major package has a direct role in fulfilling the hackathon requirements, "
        "and the documentation set is structured so judges can understand both how to run the system and why the architecture makes sense."
    )

    doc.save(path)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    build_stack_report(DOCS_DIR / "Stack Report & Justification.docx")
    build_dependencies_doc(DOCS_DIR / "Dependencies & Documentation.docx")


if __name__ == "__main__":
    main()
